import json
from collections import defaultdict
from datetime import time as dtime, timedelta

from django.contrib import messages
from django.http import HttpResponse, JsonResponse
from django.shortcuts import get_object_or_404, redirect, render
from django.template.loader import render_to_string
from django.utils import timezone
from django.views import View

from apps.core.mixins import CommitteeRequiredMixin, OrganizerRequiredMixin
from apps.submissions.models import Proposal

from .forms import AnnexEventForm, CommunicationForm, SessionForm
from .models import AnnexEvent, Communication, Session


def _event_days(event):
    """Retourne la liste des jours couverts par l'événement."""
    days = []
    current = event.start_date
    while current <= event.end_date:
        days.append(current)
        current += timedelta(days=1)
    return days


def _build_programme_context(event, sort_by_order=False):
    """Construit le contexte commun aux vues programme (organisateur et lecture).

    sort_by_order=True : sessions triées par order (vue construction D&D).
    sort_by_order=False : tout trié par start_time (vue calendrier / PDF).
    """
    days = _event_days(event)
    sessions_by_day = defaultdict(list)
    for session in (
        Session.objects.filter(event=event)
        .prefetch_related("communications__proposal")
        .order_by("date", "order", "start_time")
    ):
        sessions_by_day[session.date].append(session)

    annex_order = ("date", "order", "start_time") if sort_by_order else ("date", "start_time")
    annex_by_day = defaultdict(list)
    for item in AnnexEvent.objects.filter(event=event).order_by(*annex_order):
        annex_by_day[item.date].append(item)

    programme = []
    for day in days:
        session_items = [
            {"kind": "session", "obj": s, "start": s.start_time, "end": s.end_time}
            for s in sessions_by_day[day]
        ]
        annex_items = [
            {"kind": "annex", "obj": a, "start": a.start_time, "end": a.end_time}
            for a in annex_by_day[day]
        ]
        items = session_items + annex_items
        if not sort_by_order:
            items.sort(key=lambda x: x["start"] or dtime(0, 0))
        programme.append({
            "date": day,
            "items": items,
            "sessions": sessions_by_day[day],
            "annexes": annex_by_day[day],
        })

    return {
        "days": days,
        "programme": programme,
        "sessions_by_day": dict(sessions_by_day),
    }


# ── Vues organisateur ─────────────────────────────────────────────────────────

class ProgrammeView(OrganizerRequiredMixin, View):
    def get(self, request, event_slug):
        ctx = _build_programme_context(self.event, sort_by_order=True)
        accepted = list(
            Proposal.objects.filter(
                event=self.event,
                status=Proposal.Status.ACCEPTED,
            ).prefetch_related("authors")
        )
        scheduled_ids = set(
            Communication.objects.filter(session__event=self.event, proposal__isnull=False)
            .values_list("proposal_id", flat=True)
        )
        unscheduled = [p for p in accepted if p.pk not in scheduled_ids]
        proposals_data = {
            str(p.pk): {
                "title": p.title,
                "speaker": ", ".join(a.full_name for a in p.authors.all()),
            }
            for p in accepted
        }
        return render(request, "programme/organizer/programme.html", {
            "event": self.event,
            "membership": self.membership,
            "session_form": SessionForm(event=self.event),
            "annex_form": AnnexEventForm(event=self.event),
            "comm_form": CommunicationForm(event=self.event),
            "unscheduled_proposals": unscheduled,
            "proposals_data": proposals_data,
            **ctx,
        })


class SessionCreateView(OrganizerRequiredMixin, View):
    def post(self, request, event_slug):
        form = SessionForm(request.POST, event=self.event)
        if form.is_valid():
            session = form.save(commit=False)
            session.event = self.event
            # ordre = dernier + 1 pour ce jour
            last = (
                Session.objects.filter(event=self.event, date=session.date)
                .order_by("-order")
                .values_list("order", flat=True)
                .first()
            )
            session.order = (last or 0) + 1
            session.save()
            messages.success(request, "Session créée.")
        else:
            messages.error(request, "Erreur dans le formulaire de session.")
        return redirect("programme:programme", event_slug=event_slug)


class SessionUpdateView(OrganizerRequiredMixin, View):
    def post(self, request, event_slug, session_id):
        session = get_object_or_404(Session, pk=session_id, event=self.event)
        form = SessionForm(request.POST, instance=session, event=self.event)
        if form.is_valid():
            form.save()
            messages.success(request, "Session mise à jour.")
        return redirect("programme:programme", event_slug=event_slug)


class SessionDeleteView(OrganizerRequiredMixin, View):
    def post(self, request, event_slug, session_id):
        session = get_object_or_404(Session, pk=session_id, event=self.event)
        session.delete()
        messages.success(request, "Session supprimée.")
        return redirect("programme:programme", event_slug=event_slug)


class SessionReorderView(OrganizerRequiredMixin, View):
    """POST JSON [{id, order}] — réordonne les sessions d'un même jour."""

    def post(self, request, event_slug):
        try:
            data = json.loads(request.body)
        except (json.JSONDecodeError, ValueError):
            return JsonResponse({"error": "JSON invalide."}, status=400)
        if not isinstance(data, list) or not data:
            return JsonResponse({"error": "Format invalide."}, status=400)
        try:
            order_map = {int(item["id"]): int(item["order"]) for item in data}
        except (KeyError, ValueError, TypeError):
            return JsonResponse({"error": "Format invalide."}, status=400)

        orders = list(order_map.values())
        if len(set(orders)) != len(orders) or any(v <= 0 for v in orders):
            return JsonResponse({"error": "Ordres invalides."}, status=400)

        sessions = list(Session.objects.filter(event=self.event, pk__in=order_map.keys()))
        if set(s.pk for s in sessions) != set(order_map.keys()):
            return JsonResponse({"error": "Sessions introuvables."}, status=400)

        now = timezone.now()
        for session in sessions:
            session.order = order_map[session.pk]
            session.updated_at = now
        Session.objects.bulk_update(sessions, ["order", "updated_at"])
        return JsonResponse({"ok": True})


class CommunicationCreateView(OrganizerRequiredMixin, View):
    def post(self, request, event_slug, session_id):
        session = get_object_or_404(Session, pk=session_id, event=self.event)
        form = CommunicationForm(request.POST, event=self.event)
        if form.is_valid():
            comm = form.save(commit=False)
            comm.session = session
            last = (
                Communication.objects.filter(session=session)
                .order_by("-order")
                .values_list("order", flat=True)
                .first()
            )
            comm.order = (last or 0) + 1
            comm.save()
            messages.success(request, "Communication ajoutée.")
        else:
            messages.error(request, "Erreur dans le formulaire.")
        return redirect("programme:programme", event_slug=event_slug)


class CommunicationUpdateView(OrganizerRequiredMixin, View):
    def post(self, request, event_slug, comm_id):
        comm = get_object_or_404(Communication, pk=comm_id, session__event=self.event)
        form = CommunicationForm(request.POST, instance=comm, event=self.event)
        if form.is_valid():
            form.save()
            messages.success(request, "Communication mise à jour.")
        return redirect("programme:programme", event_slug=event_slug)


class CommunicationDeleteView(OrganizerRequiredMixin, View):
    def post(self, request, event_slug, comm_id):
        comm = get_object_or_404(Communication, pk=comm_id, session__event=self.event)
        comm.delete()
        messages.success(request, "Communication supprimée.")
        return redirect("programme:programme", event_slug=event_slug)


class CommunicationReorderView(OrganizerRequiredMixin, View):
    """POST JSON [{id, order, session_id}] — réordonne et déplace entre sessions."""

    def post(self, request, event_slug):
        try:
            data = json.loads(request.body)
        except (json.JSONDecodeError, ValueError):
            return JsonResponse({"error": "JSON invalide."}, status=400)
        if not isinstance(data, list) or not data:
            return JsonResponse({"error": "Format invalide."}, status=400)
        try:
            updates = [
                {"id": int(item["id"]), "order": int(item["order"]), "session_id": int(item["session_id"])}
                for item in data
            ]
        except (KeyError, ValueError, TypeError):
            return JsonResponse({"error": "Format invalide."}, status=400)

        comm_ids = [u["id"] for u in updates]
        session_ids = {u["session_id"] for u in updates}

        # Vérifie que toutes les sessions appartiennent à cet événement
        valid_sessions = set(
            Session.objects.filter(event=self.event, pk__in=session_ids).values_list("pk", flat=True)
        )
        if valid_sessions != session_ids:
            return JsonResponse({"error": "Sessions introuvables."}, status=400)

        comms = {c.pk: c for c in Communication.objects.filter(pk__in=comm_ids, session__event=self.event)}
        if set(comms.keys()) != set(comm_ids):
            return JsonResponse({"error": "Communications introuvables."}, status=400)

        # Vérifie l'unicité des ordres par session
        orders_by_session = defaultdict(list)
        for u in updates:
            orders_by_session[u["session_id"]].append(u["order"])
        for session_orders in orders_by_session.values():
            if len(set(session_orders)) != len(session_orders) or any(v <= 0 for v in session_orders):
                return JsonResponse({"error": "Ordres invalides dans une session."}, status=400)

        for u in updates:
            comm = comms[u["id"]]
            comm.session_id = u["session_id"]
            comm.order = u["order"]
        Communication.objects.bulk_update(list(comms.values()), ["session_id", "order"])
        return JsonResponse({"ok": True})


class AnnexReorderView(OrganizerRequiredMixin, View):
    """POST JSON [{id, order}] — réordonne les annexes d'un même jour."""

    def post(self, request, event_slug):
        try:
            data = json.loads(request.body)
        except (json.JSONDecodeError, ValueError):
            return JsonResponse({"error": "JSON invalide."}, status=400)
        if not isinstance(data, list) or not data:
            return JsonResponse({"error": "Format invalide."}, status=400)
        try:
            order_map = {int(item["id"]): int(item["order"]) for item in data}
        except (KeyError, ValueError, TypeError):
            return JsonResponse({"error": "Format invalide."}, status=400)

        orders = list(order_map.values())
        if len(set(orders)) != len(orders) or any(v <= 0 for v in orders):
            return JsonResponse({"error": "Ordres invalides."}, status=400)

        annexes = list(AnnexEvent.objects.filter(event=self.event, pk__in=order_map.keys()))
        if set(a.pk for a in annexes) != set(order_map.keys()):
            return JsonResponse({"error": "Annexes introuvables."}, status=400)

        now = timezone.now()
        for annex in annexes:
            annex.order = order_map[annex.pk]
            annex.updated_at = now
        AnnexEvent.objects.bulk_update(annexes, ["order", "updated_at"])
        return JsonResponse({"ok": True})


class AnnexEventCreateView(OrganizerRequiredMixin, View):
    def post(self, request, event_slug):
        form = AnnexEventForm(request.POST, event=self.event)
        if form.is_valid():
            annex = form.save(commit=False)
            annex.event = self.event
            last = (
                AnnexEvent.objects.filter(event=self.event, date=annex.date)
                .order_by("-order")
                .values_list("order", flat=True)
                .first()
            )
            annex.order = (last or 0) + 1
            annex.save()
            messages.success(request, "Événement annexe ajouté.")
        else:
            messages.error(request, "Erreur dans le formulaire.")
        return redirect("programme:programme", event_slug=event_slug)


class AnnexEventUpdateView(OrganizerRequiredMixin, View):
    def post(self, request, event_slug, annex_id):
        annex = get_object_or_404(AnnexEvent, pk=annex_id, event=self.event)
        form = AnnexEventForm(request.POST, instance=annex, event=self.event)
        if form.is_valid():
            form.save()
            messages.success(request, "Événement annexe mis à jour.")
        return redirect("programme:programme", event_slug=event_slug)


class AnnexEventDeleteView(OrganizerRequiredMixin, View):
    def post(self, request, event_slug, annex_id):
        annex = get_object_or_404(AnnexEvent, pk=annex_id, event=self.event)
        annex.delete()
        messages.success(request, "Événement annexe supprimé.")
        return redirect("programme:programme", event_slug=event_slug)


# ── Export texte / Word ───────────────────────────────────────────────────────

def _fmt_date(d):
    """Formate une date sans zéro initial cross-platform (%-d non portable sur macOS/BSD)."""
    return f"{d.day} {d.strftime('%B %Y')}"


def _event_header_str(event):
    date_str = _fmt_date(event.start_date)
    if event.end_date != event.start_date:
        date_str += f" – {_fmt_date(event.end_date)}"
    if event.location:
        date_str += f" · {event.location}"
    return date_str


class ProgrammeTextExportView(OrganizerRequiredMixin, View):
    """Export du programme en .txt (universel) ou .docx (Word). ?format=txt|docx"""

    def get(self, request, event_slug):
        fmt = request.GET.get("format", "txt")
        ctx = _build_programme_context(self.event)

        if fmt == "docx":
            return self._docx(ctx)
        return self._txt(ctx)

    def _txt(self, ctx):
        lines = []
        event = self.event
        lines.append(event.name.upper())
        lines.append(_event_header_str(event))
        lines.append("")

        for day_data in ctx["programme"]:
            day_label = f"{day_data['date'].strftime('%A')} {_fmt_date(day_data['date'])}".capitalize()
            sep = "=" * len(day_label)
            lines.append(sep)
            lines.append(day_label)
            lines.append(sep)
            lines.append("")

            for item in day_data["items"]:
                if item["kind"] == "session":
                    s = item["obj"]
                    time_str = ""
                    if s.start_time:
                        time_str = f"{s.start_time.strftime('%H:%M')} – {s.end_time.strftime('%H:%M')}  "
                    loc_str = f"  [{s.location}]" if s.location else ""
                    lines.append(f"{time_str}{s.title or 'Session'}{loc_str}")
                    if s.moderator:
                        lines.append(f"  Modération : {s.moderator}")
                    for comm in s.communications.all():
                        cancelled = comm.proposal and comm.proposal.status == "cancelled"
                        kind_str = f"[{comm.get_kind_display()}] " if comm.kind != "talk" else ""
                        name_str = f"{comm.speaker_name} — " if comm.speaker_name else ""
                        annul = "  [ANNULÉ]" if cancelled else ""
                        lines.append(f"  • {kind_str}{name_str}{comm.title} ({comm.duration} min){annul}")
                    lines.append("")
                else:
                    a = item["obj"]
                    lines.append(f"{a.start_time.strftime('%H:%M')} – {a.end_time.strftime('%H:%M')}  {a.label}  ({a.get_kind_display()})")
                    lines.append("")

        content = "\n".join(lines)
        resp = HttpResponse(content, content_type="text/plain; charset=utf-8")
        resp["Content-Disposition"] = (
            f'attachment; filename="programme-{self.event.slug}.txt"'
        )
        return resp

    def _docx(self, ctx):
        from io import BytesIO
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Titre événement
        title_p = doc.add_heading(self.event.name, level=0)
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Date / lieu
        event = self.event
        meta = doc.add_paragraph(_event_header_str(event))
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta.runs[0].font.size = Pt(11)
        doc.add_paragraph("")

        for day_data in ctx["programme"]:
            day_label = f"{day_data['date'].strftime('%A')} {_fmt_date(day_data['date'])}".capitalize()
            doc.add_heading(day_label, level=1)

            for item in day_data["items"]:
                if item["kind"] == "session":
                    s = item["obj"]
                    # En-tête session
                    p = doc.add_paragraph(style="Heading 2")
                    if s.start_time:
                        run = p.add_run(f"{s.start_time.strftime('%H:%M')} – {s.end_time.strftime('%H:%M')}  ")
                        run.font.color.rgb = RGBColor(0x6C, 0x75, 0x7D)
                    p.add_run(s.title or "Session")
                    if s.location:
                        run = p.add_run(f"  · {s.location}")
                        run.font.color.rgb = RGBColor(0x6C, 0x75, 0x7D)
                    if s.moderator:
                        mod = doc.add_paragraph(f"Modération : {s.moderator}")
                        mod.runs[0].font.italic = True
                        mod.runs[0].font.size = Pt(10)
                    # Communications
                    for comm in s.communications.all():
                        cancelled = comm.proposal and comm.proposal.status == "cancelled"
                        p = doc.add_paragraph(style="List Bullet")
                        if comm.kind != "talk":
                            run = p.add_run(f"[{comm.get_kind_display()}] ")
                            run.font.italic = True
                        if comm.speaker_name:
                            run = p.add_run(f"{comm.speaker_name}")
                            run.bold = True
                            if cancelled:
                                run.font.strike = True
                            p.add_run(" — ")
                        run = p.add_run(comm.title)
                        if cancelled:
                            run.font.strike = True
                        dur = p.add_run(f"  ({comm.duration} min)")
                        dur.font.color.rgb = RGBColor(0x6C, 0x75, 0x7D)
                        if cancelled:
                            annul = p.add_run("  [ANNULÉ]")
                            annul.font.color.rgb = RGBColor(0xC2, 0x41, 0x0C)
                else:
                    a = item["obj"]
                    p = doc.add_paragraph()
                    run = p.add_run(f"{a.start_time.strftime('%H:%M')} – {a.end_time.strftime('%H:%M')}  ")
                    run.font.color.rgb = RGBColor(0x6C, 0x75, 0x7D)
                    run = p.add_run(f"{a.label}")
                    run.bold = True
                    p.add_run(f"  ({a.get_kind_display()})")

            doc.add_paragraph("")

        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        resp = HttpResponse(
            buf.read(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        resp["Content-Disposition"] = (
            f'attachment; filename="programme-{self.event.slug}.docx"'
        )
        return resp


# ── Export PDF ───────────────────────────────────────────────────────────────

class ProgrammePdfView(CommitteeRequiredMixin, View):
    def get(self, request, event_slug):
        from weasyprint import HTML
        ctx = _build_programme_context(self.event)
        html = render_to_string(
            "programme/programme_pdf.html",
            {"event": self.event, **ctx},
        )
        pdf = HTML(string=html, base_url=request.build_absolute_uri("/")).write_pdf()
        response = HttpResponse(pdf, content_type="application/pdf")
        response["Content-Disposition"] = (
            f'attachment; filename="programme-{self.event.slug}.pdf"'
        )
        return response


# ── Vue calendrier (lecture) ──────────────────────────────────────────────────

class CalendarView(CommitteeRequiredMixin, View):
    """Grille calendrier CSS pure — vue édition pour l'organisateur,
    lecture pour les autres membres."""

    def get(self, request, event_slug):
        ctx = _build_programme_context(self.event)
        # grid_start/grid_end calculés depuis les données déjà chargées dans ctx
        times = [
            item[key]
            for day_data in ctx["programme"]
            for item in day_data["items"]
            for key in ("start", "end")
            if item[key] is not None
        ]
        grid_start = min(times).replace(minute=0, second=0) if times else None
        grid_end = max(times) if times else None
        return render(request, "programme/calendar.html", {
            "event": self.event,
            "membership": self.membership,
            "grid_start": grid_start,
            "grid_end": grid_end,
            **ctx,
        })
